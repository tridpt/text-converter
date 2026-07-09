"""Document family converters. Hub representation: an HTML string.

Supported formats: txt, md, html, docx, pdf.
"""

from __future__ import annotations

import base64
import html as html_lib
import io

import mammoth
import markdown as md_lib
import pdfplumber
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches
from markdownify import markdownify
from odf import text as odf_text
from odf.opendocument import OpenDocumentText
from odf.opendocument import load as odf_load
from odf.teletype import extractText
from striprtf.striprtf import rtf_to_text
from xhtml2pdf import pisa

from ..config import settings
from .registry import (
    ConversionError,
    ConvertOptions,
    FormatSpec,
    reader,
    register_format,
    set_html_sanitizer,
    set_toc_transformer,
    writer,
)

# --- Format catalogue -------------------------------------------------------
register_format(FormatSpec("txt", "document", "Plain text", ".txt", "text/plain", False))
register_format(FormatSpec("md", "document", "Markdown", ".md", "text/markdown", False))
register_format(FormatSpec("html", "document", "HTML", ".html", "text/html", False))
register_format(
    FormatSpec(
        "docx",
        "document",
        "Word (DOCX)",
        ".docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        True,
    )
)
register_format(FormatSpec("pdf", "document", "PDF", ".pdf", "application/pdf", True))
register_format(
    FormatSpec("rtf", "document", "Rich Text (RTF)", ".rtf", "application/rtf", False)
)
register_format(
    FormatSpec(
        "odt",
        "document",
        "OpenDocument (ODT)",
        ".odt",
        "application/vnd.oasis.opendocument.text",
        True,
    )
)
register_format(
    FormatSpec("latex", "document", "LaTeX", ".tex", "application/x-tex", False)
)
# The following formats are handled by Pandoc only (see pandoc_ext.py). They
# appear as readable/writable in the catalogue only when Pandoc is available.
register_format(
    FormatSpec("epub", "document", "EPUB (ebook)", ".epub", "application/epub+zip", True)
)
register_format(
    FormatSpec("revealjs", "document", "Reveal.js slides", ".html", "text/html", False)
)
register_format(
    FormatSpec(
        "pptx",
        "document",
        "PowerPoint slides (PPTX)",
        ".pptx",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        True,
    )
)


# --- Helpers ----------------------------------------------------------------
_BLOCK_TAGS = {"p", "div", "section", "article", "br", "li", "tr", "pre", "blockquote"}
_HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}


def _decode(data: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


# CSS themes applied to HTML/PDF output.
THEMES = {
    "default": (
        "body{font-family:sans-serif;line-height:1.5;max-width:800px;"
        "margin:2rem auto;padding:0 1rem;color:#222;}"
        "pre{background:#f4f4f4;padding:1rem;overflow:auto;}"
        "table{border-collapse:collapse;}td,th{border:1px solid #ccc;padding:6px;}"
    ),
    "github": (
        "body{font-family:-apple-system,Segoe UI,Helvetica,Arial,sans-serif;"
        "line-height:1.6;max-width:820px;margin:2rem auto;padding:0 1rem;color:#24292f;}"
        "h1,h2{border-bottom:1px solid #d0d7de;padding-bottom:.3em;}"
        "code,pre{background:#f6f8fa;border-radius:6px;}"
        "pre{padding:1rem;overflow:auto;}"
        "table{border-collapse:collapse;}td,th{border:1px solid #d0d7de;padding:6px 13px;}"
        "blockquote{color:#57606a;border-left:.25em solid #d0d7de;padding:0 1em;}"
    ),
    "dark": (
        "body{font-family:sans-serif;line-height:1.6;max-width:800px;"
        "margin:2rem auto;padding:1rem;background:#0d1117;color:#c9d1d9;}"
        "a{color:#58a6ff;}pre,code{background:#161b22;}"
        "pre{padding:1rem;overflow:auto;border-radius:6px;}"
        "table{border-collapse:collapse;}td,th{border:1px solid #30363d;padding:6px;}"
    ),
    "minimal": (
        "body{font-family:Georgia,serif;line-height:1.7;max-width:700px;"
        "margin:3rem auto;padding:0 1rem;color:#111;}"
    ),
}

_PAPER_SIZES = {"A3", "A4", "A5", "Letter", "Legal"}


def _wrap_html_document(
    body_html: str, theme: str = "default", paper_size: str | None = None
) -> str:
    """Wrap fragment HTML in a well-formed HTML document with a theme."""
    css = THEMES.get(theme, THEMES["default"])
    if paper_size and paper_size in _PAPER_SIZES:
        css += f"@page{{size:{paper_size};margin:2cm;}}"
    css += (
        ".toc{border:1px solid #ccc;padding:.5rem 1rem;margin-bottom:1.5rem;}"
        ".toc ul{margin:.3rem 0;}"
    )
    return (
        "<!DOCTYPE html>\n<html>\n<head>\n"
        '<meta charset="utf-8">\n'
        f"<style>{css}</style>\n</head>\n<body>\n" + body_html + "\n</body>\n</html>\n"
    )


def _slugify(text: str, used: set[str]) -> str:
    base = "".join(c if c.isalnum() else "-" for c in text.lower()).strip("-")
    base = base or "section"
    slug = base
    i = 1
    while slug in used:
        i += 1
        slug = f"{base}-{i}"
    used.add(slug)
    return slug


def add_table_of_contents(html: str) -> str:
    """Insert anchor ids on headings and prepend a nested table of contents."""
    soup = BeautifulSoup(html, "html.parser")
    headings = soup.find_all(_HEADING_TAGS)
    if not headings:
        return html

    used_ids: set[str] = set()
    entries = []
    for h in headings:
        text = h.get_text(strip=True)
        if not text:
            continue
        anchor = h.get("id") or _slugify(text, used_ids)
        h["id"] = anchor
        entries.append((int(h.name[1]), text, anchor))

    if not entries:
        return html

    min_level = min(level for level, _, _ in entries)
    items = []
    for level, text, anchor in entries:
        indent = f"margin-left:{(level - min_level) * 20}px"
        items.append(
            f'<li style="{indent}"><a href="#{anchor}">{html_lib.escape(text)}</a></li>'
        )
    toc = (
        '<div class="toc"><strong>Contents</strong><ul>' + "".join(items) + "</ul></div>"
    )

    target = soup.body or soup
    first = target.find(True)
    if first is not None:
        first.insert_before(BeautifulSoup(toc, "html.parser"))
    else:
        target.append(BeautifulSoup(toc, "html.parser"))
    return str(soup)


# Register TOC transformer with the registry (used when options.toc is set).
set_toc_transformer(add_table_of_contents)


# --- HTML sanitization (defense against malicious/active content) -----------
_DANGEROUS_TAGS = {"script", "iframe", "object", "embed", "applet"}


def sanitize_html(html: str) -> str:
    """Strip scripts and other active content from an HTML fragment/document."""
    if not settings.sanitize_html:
        return html
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup.find_all(_DANGEROUS_TAGS):
        tag.decompose()
    for el in soup.find_all(True):
        for attr in list(el.attrs):
            low = attr.lower()
            if low.startswith("on"):  # event handlers: onclick, onerror, ...
                del el[attr]
            elif low in ("href", "src", "xlink:href"):
                value = str(el.get(attr, "")).strip().lower()
                if value.startswith("javascript:") or value.startswith("vbscript:"):
                    del el[attr]
                elif value.startswith("data:") and not value.startswith("data:image/"):
                    del el[attr]
    return str(soup)


set_html_sanitizer(sanitize_html)


# --- Readers (bytes -> HTML string) -----------------------------------------
@reader("txt")
def read_txt(data: bytes) -> str:
    text = _decode(data)
    paragraphs = text.split("\n\n")
    parts = []
    for para in paragraphs:
        para = para.strip("\n")
        if not para:
            continue
        escaped = html_lib.escape(para).replace("\n", "<br>\n")
        parts.append(f"<p>{escaped}</p>")
    return "\n".join(parts)


@reader("md")
def read_md(data: bytes) -> str:
    text = _decode(data)
    return md_lib.markdown(text, extensions=["extra", "sane_lists", "tables"])


@reader("html")
def read_html(data: bytes) -> str:
    return _decode(data)


@reader("docx")
def read_docx(data: bytes) -> str:
    # Embed images inline as base64 data URIs so they survive later conversions.
    result = mammoth.convert_to_html(
        io.BytesIO(data), convert_image=mammoth.images.data_uri
    )
    return result.value


def _extract_pdf_images(page) -> list[str]:
    """Return a list of data-URI <img> tags for images embedded in a pdf page."""
    tags = []
    try:
        from pypdf import PdfReader  # noqa: F401  (import kept local/optional)
    except Exception:  # noqa: BLE001
        return tags
    for image in getattr(page, "images", []) or []:
        try:
            raw = image.data
            mime = "image/png"
            name = (getattr(image, "name", "") or "").lower()
            if name.endswith((".jpg", ".jpeg")):
                mime = "image/jpeg"
            elif name.endswith(".gif"):
                mime = "image/gif"
            encoded = base64.b64encode(raw).decode("ascii")
            tags.append(f'<img src="data:{mime};base64,{encoded}">')
        except Exception:  # noqa: BLE001 - skip images we cannot decode
            continue
    return tags


@reader("pdf")
def read_pdf(data: bytes) -> str:
    parts = []
    # Text via pdfplumber (good layout handling).
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            for para in text.split("\n"):
                para = para.strip()
                if para:
                    parts.append(f"<p>{html_lib.escape(para)}</p>")

    # Images via pypdf (extracted as inline data URIs).
    try:
        from pypdf import PdfReader

        reader_obj = PdfReader(io.BytesIO(data))
        for page in reader_obj.pages:
            parts.extend(_extract_pdf_images(page))
    except Exception:  # noqa: BLE001 - image extraction is best-effort
        pass

    return "\n".join(parts)


# --- Writers (HTML string -> bytes) -----------------------------------------
@writer("html")
def write_html(html: str, options: ConvertOptions | None = None) -> bytes:
    options = options or ConvertOptions()
    # Only wrap if it is a bare fragment.
    if "<html" not in html.lower():
        html = _wrap_html_document(html, theme=options.theme)
    return html.encode("utf-8")


@writer("md")
def write_md(html: str, options: ConvertOptions | None = None) -> bytes:
    md_text = markdownify(html, heading_style="ATX")
    # Collapse excessive blank lines.
    lines = [ln.rstrip() for ln in md_text.splitlines()]
    cleaned = "\n".join(lines).strip() + "\n"
    return cleaned.encode("utf-8")


@writer("txt")
def write_txt(html: str, options: ConvertOptions | None = None) -> bytes:
    soup = BeautifulSoup(html, "html.parser")
    for br in soup.find_all("br"):
        br.replace_with("\n")
    blocks = soup.find_all(_BLOCK_TAGS | _HEADING_TAGS)
    if blocks:
        text = "\n\n".join(
            b.get_text(strip=True) for b in blocks if b.get_text(strip=True)
        )
    else:
        text = soup.get_text("\n")
    return (text.strip() + "\n").encode("utf-8")


def _data_uri_to_stream(src: str) -> io.BytesIO | None:
    """Decode a ``data:`` image URI into a byte stream, else return None."""
    if not src or not src.startswith("data:"):
        return None
    try:
        _, encoded = src.split(",", 1)
        return io.BytesIO(base64.b64decode(encoded))
    except (ValueError, base64.binascii.Error):
        return None


def _add_runs(paragraph, element: Tag) -> None:
    """Add text runs from ``element`` to a docx paragraph, keeping bold/italic."""
    for node in element.children:
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                paragraph.add_run(text)
        elif isinstance(node, Tag):
            if node.name == "br":
                paragraph.add_run().add_break()
                continue
            run = paragraph.add_run(node.get_text())
            if node.name in ("strong", "b"):
                run.bold = True
            elif node.name in ("em", "i"):
                run.italic = True
            elif node.name == "code":
                run.font.name = "Consolas"


def _add_table(doc, table_el: Tag) -> None:
    rows = table_el.find_all("tr")
    if not rows:
        return
    ncols = max(len(r.find_all(["td", "th"])) for r in rows)
    if ncols == 0:
        return
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    for tr in rows:
        cells = tr.find_all(["td", "th"])
        row = table.add_row().cells
        for i, cell in enumerate(cells):
            if i < ncols:
                row[i].text = cell.get_text(strip=True)


def _walk_docx(doc, element: Tag) -> None:
    """Recursively emit docx content from block-level HTML elements."""
    for child in element.children:
        if isinstance(child, NavigableString):
            continue
        if not isinstance(child, Tag):
            continue
        name = child.name
        if name in _HEADING_TAGS:
            doc.add_heading(child.get_text(strip=True), level=int(name[1]))
        elif name == "p":
            para = doc.add_paragraph()
            _add_runs(para, child)
            img = child.find("img")
            if img:
                _emit_image(doc, img)
        elif name in ("ul", "ol"):
            style = "List Bullet" if name == "ul" else "List Number"
            for li in child.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(strip=True), style=style)
        elif name == "table":
            _add_table(doc, child)
        elif name == "pre":
            doc.add_paragraph(child.get_text()).style = "No Spacing"
        elif name == "blockquote":
            doc.add_paragraph(child.get_text(strip=True), style="Intense Quote")
        elif name == "img":
            _emit_image(doc, child)
        elif name in ("div", "section", "article", "body", "main", "header", "footer"):
            _walk_docx(doc, child)
        else:
            text = child.get_text(strip=True)
            if text:
                doc.add_paragraph(text)


def _emit_image(doc, img: Tag) -> None:
    stream = _data_uri_to_stream(img.get("src", ""))
    if stream is None:
        return
    try:
        doc.add_picture(stream, width=Inches(5))
    except Exception:  # noqa: BLE001 - unsupported/broken image, skip gracefully
        pass


@writer("docx")
def write_docx(html: str, options: ConvertOptions | None = None) -> bytes:
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    root = soup.body or soup

    _walk_docx(doc, root)

    # Fallback: nothing block-level was found, dump raw text.
    if not doc.paragraphs and not doc.tables:
        for line in soup.get_text("\n").splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@writer("pdf")
def write_pdf(html: str, options: ConvertOptions | None = None) -> bytes:
    options = options or ConvertOptions()
    if "<html" not in html.lower():
        html = _wrap_html_document(
            html, theme=options.theme, paper_size=options.paper_size
        )
    buf = io.BytesIO()
    result = pisa.CreatePDF(src=html, dest=buf, encoding="utf-8")
    if result.err:
        raise ConversionError("Failed to render PDF from HTML.")
    return buf.getvalue()


# --- RTF ---------------------------------------------------------------------
def _rtf_escape(text: str) -> str:
    out = []
    for ch in text:
        code = ord(ch)
        if ch in "\\{}":
            out.append("\\" + ch)
        elif ch == "\n":
            out.append("\\line ")
        elif code > 127:
            out.append(f"\\u{code}?")
        else:
            out.append(ch)
    return "".join(out)


def _rtf_inline(element: Tag) -> str:
    parts = []
    for node in element.children:
        if isinstance(node, NavigableString):
            parts.append(_rtf_escape(str(node)))
        elif isinstance(node, Tag):
            if node.name == "br":
                parts.append("\\line ")
            elif node.name in ("strong", "b"):
                parts.append("{\\b " + _rtf_escape(node.get_text()) + "}")
            elif node.name in ("em", "i"):
                parts.append("{\\i " + _rtf_escape(node.get_text()) + "}")
            else:
                parts.append(_rtf_escape(node.get_text()))
    return "".join(parts)


@reader("rtf")
def read_rtf(data: bytes) -> str:
    text = rtf_to_text(_decode(data))
    parts = []
    for para in text.split("\n\n"):
        para = para.strip()
        if para:
            escaped = html_lib.escape(para).replace("\n", "<br>\n")
            parts.append(f"<p>{escaped}</p>")
    return "\n".join(parts)


@writer("rtf")
def write_rtf(html: str, options: ConvertOptions | None = None) -> bytes:
    soup = BeautifulSoup(html, "html.parser")
    root = soup.body or soup
    body_parts = []
    for el in root.find_all(_HEADING_TAGS | {"p", "li", "pre", "blockquote"}):
        text = el.get_text(strip=True)
        if not text:
            continue
        name = el.name
        if name in _HEADING_TAGS:
            size = {1: 36, 2: 32, 3: 28, 4: 26, 5: 24, 6: 22}.get(int(name[1]), 24)
            body_parts.append(f"{{\\b\\fs{size} {_rtf_escape(text)}}}\\par")
        elif name == "li":
            body_parts.append("\\bullet\\tab " + _rtf_inline(el) + "\\par")
        elif name == "pre":
            body_parts.append("{\\f1 " + _rtf_escape(text) + "}\\par")
        elif name == "blockquote":
            body_parts.append("{\\i " + _rtf_escape(text) + "}\\par")
        else:
            body_parts.append(_rtf_inline(el) + "\\par")
    body = "\n".join(body_parts)
    doc = (
        "{\\rtf1\\ansi\\ansicpg1252\\deff0"
        "{\\fonttbl{\\f0 Calibri;}{\\f1 Consolas;}}\n" + body + "\n}"
    )
    return doc.encode("utf-8")


# --- ODT ---------------------------------------------------------------------
@reader("odt")
def read_odt(data: bytes) -> str:
    doc = odf_load(io.BytesIO(data))
    parts = []
    for el in doc.text.childNodes:
        name = getattr(el, "qname", (None, None))[1]
        if name == "h":
            try:
                level = int(el.getAttribute("outlinelevel") or 1)
            except (TypeError, ValueError):
                level = 1
            level = min(max(level, 1), 6)
            txt = extractText(el)
            if txt.strip():
                parts.append(f"<h{level}>{html_lib.escape(txt)}</h{level}>")
        elif name == "p":
            txt = extractText(el)
            if txt.strip():
                parts.append(f"<p>{html_lib.escape(txt)}</p>")
    return "\n".join(parts)


@writer("odt")
def write_odt(html: str, options: ConvertOptions | None = None) -> bytes:
    soup = BeautifulSoup(html, "html.parser")
    root = soup.body or soup
    doc = OpenDocumentText()
    for el in root.find_all(_HEADING_TAGS | {"p", "li", "pre", "blockquote"}):
        text = el.get_text(strip=True)
        if not text:
            continue
        if el.name in _HEADING_TAGS:
            doc.text.addElement(odf_text.H(outlinelevel=int(el.name[1]), text=text))
        elif el.name == "li":
            doc.text.addElement(odf_text.P(text="\u2022 " + text))
        else:
            doc.text.addElement(odf_text.P(text=text))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- LaTeX (writer only) -----------------------------------------------------
_LATEX_SPECIAL = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def _latex_escape(text: str) -> str:
    return "".join(_LATEX_SPECIAL.get(ch, ch) for ch in text)


@writer("latex")
def write_latex(html: str, options: ConvertOptions | None = None) -> bytes:
    soup = BeautifulSoup(html, "html.parser")
    root = soup.body or soup
    body_parts = []
    section_cmds = {1: "section", 2: "subsection", 3: "subsubsection"}

    for el in root.find_all(_HEADING_TAGS | {"p", "ul", "ol", "pre", "blockquote"}):
        name = el.name
        if name in _HEADING_TAGS:
            level = int(name[1])
            cmd = section_cmds.get(level, "paragraph")
            body_parts.append(f"\\{cmd}*{{{_latex_escape(el.get_text(strip=True))}}}")
        elif name in ("ul", "ol"):
            env = "itemize" if name == "ul" else "enumerate"
            items = [
                f"  \\item {_latex_escape(li.get_text(strip=True))}"
                for li in el.find_all("li", recursive=False)
                if li.get_text(strip=True)
            ]
            if items:
                body_parts.append(
                    f"\\begin{{{env}}}\n" + "\n".join(items) + f"\n\\end{{{env}}}"
                )
        elif name == "pre":
            body_parts.append("\\begin{verbatim}\n" + el.get_text() + "\n\\end{verbatim}")
        elif name == "blockquote":
            body_parts.append(
                "\\begin{quote}\n"
                + _latex_escape(el.get_text(strip=True))
                + "\n\\end{quote}"
            )
        else:  # paragraph
            text = el.get_text(strip=True)
            if text:
                body_parts.append(_latex_escape(text))

    body = "\n\n".join(body_parts)
    document = (
        "\\documentclass{article}\n"
        "\\usepackage[utf8]{inputenc}\n"
        "\\usepackage[T1]{fontenc}\n"
        "\\usepackage{graphicx}\n"
        "\\usepackage{hyperref}\n"
        "\\begin{document}\n\n" + body + "\n\n\\end{document}\n"
    )
    return document.encode("utf-8")
